"""Generate Mov-Ment application documentation as DOCX."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(ROOT, "Mov-Ment-Application-Documentation.docx")


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("Mov-Ment\nEvent Booking & Management Application")
    t.bold = True
    t.font.size = Pt(24)
    t.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(f"Technical Documentation\nVersion 1.0.0 | {date.today().strftime('%B %d, %Y')}")
    s.font.size = Pt(12)
    s.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # 1. Overview
    add_heading(doc, "1. Application Overview", 1)
    add_para(doc, "Mov-Ment is a full-stack web application for event booking and management. Customers book events (birthdays, corporate events, farewells, etc.), managers accept and execute them, and administrators oversee the entire platform.")
    add_bullets(doc, [
        "Project name: Mov-Ment",
        "Type: Event booking & management platform",
        "Architecture: React SPA + Node.js REST API + MongoDB",
        "Development URLs: Frontend http://localhost:5173 | API http://localhost:5000",
        "Production (single-server): One URL serves UI + API (typically port 5000)",
    ])

    add_heading(doc, "Core Workflow", 2)
    add_bullets(doc, [
        "User registers/logs in and creates an event booking with location, services, and schedule.",
        "Booking enters 'pending' status; managers can accept it manually.",
        "After 15 minutes without acceptance, auto-assign assigns a manager in the same city.",
        "Manager updates status: accepted → in_progress → completed.",
        "User can cancel/reschedule (pending/accepted), download invoices, submit feedback and surveys.",
        "Admin/owner manages users, managers, analytics, refunds, promotions, and support tickets.",
    ])

    # 2. Tech Stack
    add_heading(doc, "2. Technology Stack", 1)
    add_table(doc, ["Layer", "Technologies"], [
        ["Frontend", "React 19, Vite 7, React Router 7, JavaScript (JSX), CSS"],
        ["Backend", "Node.js 18+, Express 5, Mongoose 9"],
        ["Database", "MongoDB (local / Atlas); in-memory fallback in dev only"],
        ["Authentication", "JWT (7-day expiry), bcrypt, Speakeasy TOTP (2FA)"],
        ["Security", "Helmet, express-rate-limit, CORS, compression"],
        ["PDF", "PDFKit (invoice generation)"],
        ["DevOps", "Docker, Docker Compose, GitHub Actions CI, concurrently"],
    ])

    # 3. Roles
    add_heading(doc, "3. User Roles & Permissions", 1)
    add_table(doc, ["Role", "Description", "Dashboard", "Notes"], [
        ["user", "Customer who books events", "/user", "Auto-approved on registration"],
        ["manager", "Handles assigned events", "/manager", "Must be approved before login"],
        ["admin", "Platform administrator", "/admin", "Promoted from manager"],
        ["owner", "Top-level account", "/admin", "Cannot be demoted; highest privilege"],
    ])
    add_bullets(doc, [
        "Frontend protection: ProtectedRoute component checks localStorage token and role.",
        "Backend protection: JWT auth middleware; admin routes re-verify role in database.",
        "Managers registering via /register start with approved=false until admin approves.",
    ])

    # 4. Features by role
    add_heading(doc, "4. Features by Role", 1)

    add_heading(doc, "4.1 User Dashboard (/user)", 2)
    add_bullets(doc, [
        "Dashboard: upcoming events, package suggestions with past examples.",
        "Book event: type, title, date/time, address, guest count, venue, additional services, custom requests.",
        "My events: history, cancel, reschedule, invoice download (PDF/JSON), feedback, survey.",
        "Notifications: view, mark read (individual or all).",
        "Chat with manager: per-event conversations after manager assignment.",
        "Support: create tickets (query/complaint/feedback), view and reply.",
        "Payments: record payments (stub), view history and receipts.",
        "FAQ: static help content.",
        "Profile: update name, phone, picture, location.",
        "Security: enable/disable 2FA; request manager role upgrade.",
    ])

    add_heading(doc, "4.2 Manager Dashboard (/manager)", 2)
    add_bullets(doc, [
        "My events: filter by date, type, city, status; update status; assign team.",
        "Pending events: view and accept bookings.",
        "Nearby events: location-based pending events (Haversine distance).",
        "Calendar: month view of assigned events.",
        "Performed events: upload portfolio events shown to users as examples.",
        "Customer chat: start/view conversations, send messages.",
        "Resources: CRUD inventory (decoration, equipment, catering, etc.).",
        "Performance: completion rate, status breakdown, feedback with reply.",
        "Notifications and event reminders.",
        "Profile and GPS location update for nearby matching.",
    ])

    add_heading(doc, "4.3 Admin Dashboard (/admin)", 2)
    add_bullets(doc, [
        "All users: list, set/reset passwords.",
        "Manager requests: approve/reject user upgrade requests.",
        "Pending managers: approve newly registered managers.",
        "Managers: list, promote to admin, remove (demote to user).",
        "Events: overview (summary), assign teams.",
        "Manager–customer chat: read-only oversight.",
        "Support tickets: view, filter, reply, update status.",
        "Refunds: create, approve, process, reject.",
        "User activity: audit log (login, bookings, etc.).",
        "Analytics: revenue, bookings, ratings, event types, manager performance, load balancing.",
        "Notifications: broadcast or targeted alerts.",
        "Promotions: create/manage discount codes.",
    ])

    add_heading(doc, "4.4 Background Services", 2)
    add_bullets(doc, [
        "Auto-assign scheduler: every 60 seconds, assigns overdue pending events to city-matched manager.",
        "Reminder scheduler: every 6 hours, sends 24-hour event reminders to users.",
    ])

    # 5. API Routes
    add_heading(doc, "5. API Routes Summary", 1)
    add_para(doc, "Base URL: /api | Health: GET /api/health | Auth header: Authorization: Bearer <token>")

    add_heading(doc, "5.1 Authentication (/api/auth)", 2)
    add_table(doc, ["Method", "Route", "Access", "Description"], [
        ["POST", "/register", "Public", "Register as user or manager"],
        ["POST", "/login", "Public", "Login with email/phone + password; 2FA challenge if enabled"],
        ["POST", "/login/verify-2fa", "Public", "Complete login with TOTP code"],
        ["POST", "/2fa/enable", "Authenticated", "Generate TOTP secret and QR URL"],
        ["POST", "/2fa/verify", "Authenticated", "Confirm 2FA setup"],
        ["POST", "/2fa/disable", "Authenticated", "Disable 2FA (requires password)"],
    ])

    add_heading(doc, "5.2 Events (/api/events)", 2)
    add_table(doc, ["Method", "Route", "Roles", "Description"], [
        ["POST", "/", "user, admin, owner", "Create event booking"],
        ["GET", "/meta", "Public", "Event types and additional services"],
        ["GET", "/suggestions", "user, admin, owner", "Package suggestions"],
        ["GET", "/similar", "user, manager, admin, owner", "Similar events by type"],
        ["GET", "/my", "Authenticated", "List own booked events"],
        ["GET", "/pending", "manager, admin, owner", "Pending events (city/distance filter)"],
        ["POST", "/:id/accept", "manager, admin, owner", "Accept pending event"],
        ["POST", "/:id/cancel", "user, admin, owner", "Cancel own event"],
        ["POST", "/:id/reschedule", "user, admin, owner", "Reschedule own event"],
        ["POST", "/:id/status", "manager, admin, owner", "Update event status"],
    ])

    add_heading(doc, "5.3 User, Manager & Admin Routes", 2)
    add_bullets(doc, [
        "/api/user — profile, notifications, support, payments, invoices, feedback, surveys, chat, FAQ, manager requests.",
        "/api/manager — assigned events, calendar, nearby, team assignment, chat, resources, performance, feedback, portfolio.",
        "/api/admin — users, managers, events, conversations, tickets, refunds, activity, analytics, notifications, promotions.",
        "All admin routes require admin or owner role verified against the database.",
    ])

    # 6. Database
    add_heading(doc, "6. Database Models", 1)
    add_table(doc, ["Model", "Purpose", "Key Relationships"], [
        ["User", "Accounts (user/manager/admin/owner)", "Referenced by all models"],
        ["Event", "Event bookings", "bookedBy → User, assignedManager → User"],
        ["ManagerPortfolioEvent", "Manager showcase portfolio", "manager → User"],
        ["ManagerConversation", "Per-event chat", "event, user, manager → User/Event"],
        ["Notification", "User alerts", "user → User, relatedEvent → Event"],
        ["Payment", "Payment records", "user → User, event → Event"],
        ["Refund", "Refund requests", "event, user, payment → related models"],
        ["SupportTicket", "Support queries", "user → User, relatedEvent → Event"],
        ["Feedback", "Event ratings", "event, user, manager → User/Event"],
        ["Survey", "Post-event surveys", "event, user → Event/User"],
        ["ManagerRequest", "Role upgrade requests", "user → User"],
        ["Resource", "Manager inventory", "manager → User"],
        ["Promotion", "Discount codes", "createdBy → User"],
        ["UserActivity", "Audit log", "user → User"],
    ])

    add_heading(doc, "Event Status Flow", 2)
    add_para(doc, "pending → accepted → in_progress → completed  |  cancelled (from pending or accepted)")

    add_heading(doc, "Event Types", 2)
    add_para(doc, "birthday, surprise, anniversary, farewell, software_launch, corporate, other")

    # 7. Environment Variables
    add_heading(doc, "7. Environment Variables", 1)
    add_heading(doc, "Development (server/.env)", 2)
    add_table(doc, ["Variable", "Example", "Description"], [
        ["NODE_ENV", "development", "Environment mode"],
        ["PORT", "5000", "API server port"],
        ["MONGO_URI", "mongodb://localhost:27017/movment", "MongoDB connection"],
        ["JWT_SECRET", "32+ character string", "JWT signing secret"],
        ["SEED_DEMO_ACCOUNTS", "true", "Auto-create test accounts"],
        ["ALLOW_MEMORY_DB", "true", "In-memory MongoDB fallback in dev"],
        ["SERVE_CLIENT", "false", "Serve React build from API"],
    ])

    add_heading(doc, "Production (server/.env)", 2)
    add_table(doc, ["Variable", "Required", "Description"], [
        ["NODE_ENV", "Yes", "Must be production"],
        ["MONGO_URI", "Yes", "MongoDB Atlas connection string"],
        ["JWT_SECRET", "Yes", "Random string, min 32 characters"],
        ["CLIENT_URL", "Yes*", "Frontend URL(s) for CORS"],
        ["SERVE_CLIENT", "Recommended", "true for single-server deploy"],
        ["SEED_DEMO_ACCOUNTS", "No", "Must be false in production"],
        ["OWNER_EMAIL", "Recommended", "Owner account email"],
        ["OWNER_PASSWORD", "Recommended", "Strong password (12+ chars)"],
        ["TRUST_PROXY", "Yes", "true behind nginx/Railway/Render"],
    ])
    add_para(doc, "*Not required when UI and API share same origin (SERVE_CLIENT=true).")

    add_heading(doc, "Client (client/.env)", 2)
    add_table(doc, ["Variable", "Description"], [
        ["VITE_API_URL", "API base URL. Empty = relative /api paths (same-origin deploy). Set for split hosting."],
    ])

    # 8. Folder Structure
    add_heading(doc, "8. Project Folder Structure", 1)
    structure = """
mov-ment/
├── README.md, OPERATIONS.md, DEPLOYMENT.md, ATLAS_SETUP.md
├── Mov-Ment-Application-Documentation.docx (this document)
├── package.json, Dockerfile, docker-compose.yml
├── scripts/          — setup-env.js, generate-documentation.py
├── client/           — React frontend (Vite)
│   ├── src/pages/    — Landing, Login, Register, Dashboards
│   ├── src/components/ — ProtectedRoute, Modal, ProfileDropdown
│   ├── src/utils/api.js — API URL helper
│   └── dist/         — Production build output
└── server/           — Express API
    ├── index.js      — Entry point
    ├── config/env.js — Environment config
    ├── middleware/   — auth, security, errorHandler, serveClient
    ├── models/       — 14 Mongoose schemas
    ├── routes/       — auth, events, user, manager, admin
    ├── services/     — database, seed, autoAssign
    └── scripts/      — createOwner, resetPassword
"""
    add_para(doc, structure)

    # 9. Install & Run
    add_heading(doc, "9. Installation & Running", 1)
    add_heading(doc, "First-time setup", 2)
    add_bullets(doc, [
        "cd E:\\Projects\\mov-ment",
        "npm run setup — installs dependencies and creates .env files",
        "Edit server/.env: set MONGO_URI and JWT_SECRET (32+ chars)",
    ])

    add_heading(doc, "Development", 2)
    add_bullets(doc, [
        "npm start — runs API (port 5000) + Vite dev server (port 5173)",
        "npm run start:server — API only",
        "npm run start:client — frontend only",
        "Stop with Ctrl + C",
    ])

    add_heading(doc, "Production", 2)
    add_bullets(doc, [
        "npm run setup:env:prod — create production .env template",
        "npm run start:prod — build frontend + run API in production mode",
        "docker compose up --build — Docker deployment",
    ])

    add_heading(doc, "Other commands", 2)
    add_table(doc, ["Command", "Description"], [
        ["npm run lint", "ESLint check (client)"],
        ["npm run build", "Production frontend build"],
        ["npm run preview", "Preview production build"],
    ])

    # 10. Security
    add_heading(doc, "10. Security Features", 1)
    add_table(doc, ["Feature", "Implementation"], [
        ["JWT Authentication", "Bearer tokens, 7-day expiry, signed with JWT_SECRET"],
        ["Password Hashing", "bcrypt with salt rounds 10"],
        ["Two-Factor Auth (2FA)", "TOTP via Speakeasy"],
        ["Role-Based Access", "Backend middleware + frontend ProtectedRoute"],
        ["Rate Limiting", "API: 300/15min (prod); Login: 20/15min (prod)"],
        ["CORS", "Restricted to CLIENT_URL in production"],
        ["Helmet", "Security headers; CSP in production"],
        ["Compression", "Enabled in production"],
        ["Error Sanitization", "No stack traces in production responses"],
        ["Manager Approval", "Unapproved managers cannot log in"],
        ["Production Validation", "Server exits on invalid production config"],
        ["Graceful Shutdown", "SIGTERM/SIGINT handlers"],
    ])

    # 11. Test Accounts
    add_heading(doc, "11. Test Accounts (Development)", 1)
    add_table(doc, ["Role", "Email", "Password", "Dashboard"], [
        ["Owner / Admin", "admin@gmail.com", "admin3168", "/admin"],
        ["User", "user@gmail.com", "user3168", "/user"],
        ["Manager", "manager@gmail.com", "manager3168", "/manager"],
    ])
    add_para(doc, "Created automatically on first server start when SEED_DEMO_ACCOUNTS=true. In production, set SEED_DEMO_ACCOUNTS=false and create owner with a strong password via OWNER_PASSWORD or: cd server && node scripts/createOwner.js")

    # 12. Deployment
    add_heading(doc, "12. Deployment Options", 1)
    add_table(doc, ["Option", "Best For", "Key Configuration"], [
        ["Single server", "Railway, Render, VPS, Docker", "SERVE_CLIENT=true, TRUST_PROXY=true"],
        ["Split hosting", "Vercel (UI) + Railway (API)", "SERVE_CLIENT=false, VITE_API_URL, CLIENT_URL"],
        ["Docker Compose", "VPS / cloud VM", "Root .env with MONGO_URI, JWT_SECRET"],
        ["VPS + nginx", "Self-hosted", "Reverse proxy + Certbot HTTPS"],
        ["PM2", "VPS process manager", "npm run build:prod then pm2 start"],
    ])

    add_heading(doc, "Pre-Launch Checklist", 2)
    add_bullets(doc, [
        "NODE_ENV=production",
        "Strong unique JWT_SECRET (32+ characters)",
        "MONGO_URI points to production Atlas cluster",
        "Atlas Network Access allows server IP",
        "SEED_DEMO_ACCOUNTS=false",
        "OWNER_PASSWORD set to strong password",
        "CLIENT_URL matches live frontend URL",
        "HTTPS enabled",
        "GET /api/health returns ok: true",
        "Login, register, and booking flow tested on live URL",
        ".env files NOT committed to Git",
    ])

    # Footer
    doc.add_page_break()
    add_heading(doc, "Document Information", 1)
    add_bullets(doc, [
        f"Generated: {date.today().strftime('%B %d, %Y')}",
        "Application: Mov-Ment v1.0.0",
        "Project path: E:\\Projects\\mov-ment",
        "Related docs: README.md, OPERATIONS.md, DEPLOYMENT.md, ATLAS_SETUP.md",
    ])

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
